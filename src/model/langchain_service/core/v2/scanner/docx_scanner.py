import docx
from docx.document import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import RGBColor

from ..element import Element


class DocxScanner:

    def __init__(self, file_obj):
        self.file_obj = file_obj
        self.doc = docx.Document(file_obj)

        self.elements = []

    def scan_elements(self):
        for block in self._iter_block_items(self.doc):
            if isinstance(block, Paragraph):
                paragraph = block
                elements = self._scan_elements_from_runs(paragraph.runs)
                self.elements.extend(elements)

            elif isinstance(block, Table):
                table = block

                prev_elements = None
                for row in table.rows:
                    for cell in row.cells:
                        elements = []
                        for paragraph in cell.paragraphs:
                            elements.extend(
                                self._scan_elements_from_runs(paragraph.runs))

                        if not self._check_if_elements_are_same(prev_elements, elements):
                            self.elements.extend(elements)
                            prev_elements = elements

    def text(self):
        text = ''

        for element in self.elements:
            text += str(element)

        return text

    def raw_text(self):
        text = ''

        for element in self.doc.paragraphs:
            text += element.text

        return text

    def _check_if_elements_are_same(self, elements1, elements2):
        if elements1 is None or elements2 is None:
            return False

        if len(elements1) != len(elements2):
            return False

        for i in range(len(elements1)):
            if elements1[i] != elements2[i]:
                return False

        return True

    def _scan_elements_from_runs(self, runs):
        elements = []
        last_element = None

        for run in runs:
            element = Element()

            element.text = run.text
            element.italic = run.italic is not None
            element.bold = run.bold is not None
            element.underline = run.underline is not None
            element.color = '#' + str(run.font.color.rgb if run.font.color.rgb is not None else RGBColor(  # type: ignore
                0, 0, 0))

            if last_element is not None and last_element.is_concatable_with(element):
                last_element.concat(element)
            else:
                elements.append(element)
                last_element = element

        newline = Element()
        newline.text = '\n'
        elements.append(newline)

        return elements

    def _iter_block_items(self, parent):
        if isinstance(parent, Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
