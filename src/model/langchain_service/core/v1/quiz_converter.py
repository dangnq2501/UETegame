import os
import docx
import uuid
import re
import zipfile
from utils import *
from .node import Node


class QuizConverter:
    # The maximum length of a question label
    # Example:
    #   "Question 1" has a label length of 2
    #   "Câu hỏi 2" has a label length of 3
    MAX_LABEL_LENGTH = 5

    def __init__(self, file_obj):
        self.file = file_obj
        self.document = docx.Document(file_obj)
        self.question_labels = []
        self.questions = []
        self.paragraphs = []
        self.id = uuid.uuid4().hex

        # mapping from relationship id to image file name
        self.rels = {}

        self._scan()
        self._extract_images()

    def get_json(self):
        """Return the JSON representation of the quiz."""

        json = {
            "testid": self.id,
            "questions": []
        }

        for question in self.questions:
            question_json = {
                "content": [],
                "choices": [[], [], [], []],
                "answer": 0
            }

            for node in question['content']:
                question_json['content'].append(node.get_json())

            for i in range(len(question['choices'])):
                for node in question['choices'][i]:
                    question_json['choices'][i].append(node.get_json())

            question_json['answer'] = question['answer']

            json['questions'].append(question_json)

        return json

    def _extract_images(self):
        """Extract all images from the document and save to disk"""

        zip_obj = zipfile.ZipFile(self.file, 'r')
        file_list = zip_obj.namelist()

        media_folder = 'word/media/'

        for file_name in file_list:
            if file_name.startswith(media_folder):
                zip_obj.extract(file_name, f'data/{self.id}/')

        zip_obj.close()

    def _fix_nodes(self, nodes):
        """Remove empty nodes and tab symbols from the list of nodes."""

        new_nodes = []
        for node in nodes:
            if node.text is None:
                new_nodes.append(node)
            elif node.text.strip() != '':
                node.text = node.text.replace('\t', '')
                new_nodes.append(node)

        return new_nodes

    def _scan(self):
        """Scan and parse the document."""

        self._scan_paragraphs()
        self._scan_label()

        current_question = None
        scanning_target = None

        for paragraph in self.paragraphs:
            i = 0

            if self._is_start_of_question(paragraph):
                current_question = {
                    'content': [],
                    'choices': [[], [], [], []],
                    'answer': None
                }
                self.questions.append(current_question)
                i = len(self.question_labels)
                scanning_target = "<content>"
                current_choice_label = None

            elif self._is_start_of_choice(paragraph):
                # start scanning choices
                node = paragraph[0]
                letter = extract_index_letter(node.text)
                scanning_target = ord(letter.lower()) - ord('a')
                i = 1
                current_choice_label = node

                if current_question == None:
                    continue

                if (node.underline and not self.choice_underline) or (node.bold and not self.choice_bold) or (
                        node.italic and not self.choice_italic) or (node.color != self.choice_color):
                    current_question["answer"] = scanning_target

            elif scanning_target != "<content>":
                continue

            if current_question == None:
                continue

            while i < len(paragraph):
                node = paragraph[i]

                if node.text is None:
                    if scanning_target == "<content>":
                        current_question['content'].append(node)
                    else:
                        current_question['choices'][scanning_target].append(
                            node)
                elif self._is_choice_label(node):
                    # start scanning choices
                    letter = extract_index_letter(node.text)
                    scanning_target = ord(letter.lower()) - ord('a')

                    if (node.underline and not self.choice_underline) or (node.bold and not self.choice_bold) or (
                            node.italic and not self.choice_italic) or (node.color != self.choice_color):
                        current_question["answer"] = scanning_target

                elif scanning_target == "<content>":
                    if len(current_question['content']) > 0 and current_question['content'][-1].is_concatable_with(
                            node):
                        current_question['content'][-1].concat(node)
                    else:
                        current_question['content'].append(node)
                else:
                    if len(current_question['choices'][scanning_target]) > 0 and \
                            current_question['choices'][scanning_target][-1].is_concatable_with(node):
                        current_question['choices'][scanning_target][-1].concat(
                            node)
                    else:
                        current_question['choices'][scanning_target].append(
                            node)

                i += 1

            if scanning_target != "<content>":
                scanning_target = None

        for question in self.questions:
            question['content'] = self._fix_nodes(question['content'])
            for j in range(len(question['choices'])):
                question['choices'][j] = self._fix_nodes(
                    question['choices'][j])

        nquestions = len(self.questions)
        print(nquestions)
        index = 0
        for table in self.document.tables:
            ncols = len(table.columns)
            nrows = len(table.rows)

            if ncols * nrows >= nquestions:
                for i in range(nrows):
                    for j in range(ncols):
                        text = table.cell(i, j).text

                        if text is None:
                            continue

                        text = text.strip()

                        if re.match(r'^\d*[ABCD]$', text) is not None and index < nquestions:
                            self.questions[index]['answer'] = ord(
                                text[-1]) - ord('A')
                            index += 1
                break

    def _scan_paragraphs(self):
        """Scan all paragraphs in the document then convert them to nodes."""

        # get relationship id of each image from the document
        for run in self.document.part.rels.values():
            if isinstance(run._target, docx.parts.image.ImagePart):  # type: ignore
                self.rels[run.rId] = os.path.basename(run._target.partname)

        for paragraph in self.document.paragraphs:
            paragraph_ = []
            skip = []
            for count, run in enumerate(paragraph.runs):
                cur = run.text
                prev = paragraph.runs[count - 1].text
                if (is_index_number(cur) is True) and (is_index_number(prev) is True):
                    paragraph.runs[count - 1].text += cur
                    skip.append(count)

            for count, run in enumerate(paragraph.runs):
                if count in skip:
                    continue
                text = run.text
                # check if the run contains an inline image
                if 'v:imagedata' in run._element.xml:
                    for key, value in self.rels.items():
                        if 'r:id="{}"'.format(key) in run._element.xml:
                            node = Node()
                            node.inline_image = value
                            paragraph_.append(node)
                            break
                # check if the run contains a drawing image
                elif 'w:drawing' in run._element.xml:
                    for key, value in self.rels.items():
                        if 'r:embed="{}"'.format(key) in run._element.xml:
                            node = Node()
                            node.image = value
                            paragraph_.append(node)
                            break
                else:
                    super_script = 'superscript' in run._element.xml
                    sub_script = 'subscript' in run._element.xml

                    subtexts = text.split(' ')
                    for j, text_ in enumerate(subtexts):
                        node = Node()
                        node.italic = run.italic is not None
                        node.bold = run.bold is not None
                        node.underline = run.underline is not None
                        node.color = run.font.color.rgb if run.font.color.rgb is not None else docx.shared.RGBColor(
                            # type: ignore
                            0, 0, 0)
                        node.super_script = super_script
                        node.sub_script = sub_script

                        node.text = text_
                        if j != len(subtexts) - 1:
                            node.text += ' '

                        paragraph_.append(node)

            self.paragraphs.append(paragraph_)

    def _scan_label(self):
        """Scan the nodes and extract the question labels."""

        # Count the number of each word in the first `MAX_LABEL_LENGTH` nodes of each paragraph
        counters = [{} for _ in range(self.MAX_LABEL_LENGTH)]

        choice_symbol_counters = {}
        choice_italic = 0
        choice_bold = 0
        choice_underline = 0
        choice_color_counters = {}

        for paragraph in self.paragraphs:
            if len(paragraph) < self.MAX_LABEL_LENGTH:
                continue

            # check choice label
            first_node = paragraph[0]
            if first_node.text is not None and is_index_letter(first_node.text):
                symbol = first_node.text.strip()[1] if len(
                    first_node.text.strip()) > 1 else ""

                if symbol not in choice_symbol_counters:
                    choice_symbol_counters[symbol] = 0

                choice_symbol_counters[symbol] += 1

                if first_node.italic:
                    choice_italic += 1
                else:
                    choice_italic -= 1

                if first_node.bold:
                    choice_bold += 1
                else:
                    choice_bold -= 1

                if first_node.underline:
                    choice_underline += 1
                else:
                    choice_underline -= 1

                if first_node.color not in choice_color_counters:
                    choice_color_counters[first_node.color] = 0

                choice_color_counters[first_node.color] += 1

            for i in range(self.MAX_LABEL_LENGTH):
                node = paragraph[i]
                name = node.text

                if node.text is None or node.text.strip() == "":
                    continue

                if is_index_number(node.text):
                    name = "<number>"
                elif is_index_letter(node.text):
                    name = "<letter>"

                if name not in counters[i]:
                    counters[i][name] = 0

                counters[i][name] += 1

        # find choice label style
        self.choice_italic = choice_italic > 0
        self.choice_bold = choice_bold > 0
        self.choice_underline = choice_underline > 0
        self.choice_symbol = max(choice_symbol_counters, key=choice_symbol_counters.get) if len(  # type: ignore
            choice_symbol_counters) > 0 else None
        self.choice_color = max(choice_color_counters, key=choice_color_counters.get) if len(  # type: ignore
            choice_color_counters) > 0 else None

        i = 0
        previous_label_count = None

        # Find the question labels
        while i < self.MAX_LABEL_LENGTH:
            total = 0

            for key, value in counters[i].items():
                if key != "<letter>":
                    total += value

            for key, value in counters[i].items():
                if key == "<letter>":
                    continue

                # current node is a question label if it has a frequency of more than 50%
                # and its number of occurrences is close to the previous node
                if value / total > 0.5 and (previous_label_count is None or is_close(value, previous_label_count, 0.2)):
                    self.question_labels.append(key)
                    previous_label_count = value
                    break

            i += 1

    def _is_start_of_question(self, paragraph):
        """Check if the given paragraph starts with question labels."""
        for word in paragraph:
            if (word is None) or (word.text == ''):
                paragraph.remove(word)

        if len(paragraph) < len(self.question_labels):
            return False

        i = 0
        while i < len(self.question_labels):
            node = paragraph[i]

            if self.question_labels[i] == "<number>":
                if not is_index_number(node.text):
                    break
            elif self.question_labels[i] != node.text:
                break

            i += 1

        return len(self.question_labels) - i <= 1

    def _is_start_of_choice(self, paragraph):
        """Check if the given paragraph starts with a choice label."""

        if len(paragraph) < 1:
            return False

        node = paragraph[0]

        if node.text is None:
            return False

        return self._is_choice_label(node)

    def _is_choice_label(self, node):
        if node.text is None or not is_index_letter(node.text):
            return False

        if len(node.text.strip()) == 1:
            if self.choice_symbol != '':
                return False
        elif node.text.strip()[1] != self.choice_symbol:
            return False

        count = 0

        if node.italic == self.choice_italic:
            count += 1

        if node.bold == self.choice_bold:
            count += 1

        if node.underline == self.choice_underline:
            count += 1

        if node.color == self.choice_color:
            count += 1

        return count >= 3
